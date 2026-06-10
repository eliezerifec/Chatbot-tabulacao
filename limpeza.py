"""
limpeza.py — Limpeza automática de bases SurveyMonkey — IFec RJ
================================================================
Fluxo:
  1. Extrai o texto do questionário em Word (.docx), onde estão descritos
     os pulos do questionário (ex.: "Se NÃO, pule para a P10").
  2. Usa IA (OpenAI) para transformar o texto em regras estruturadas,
     mapeadas para as perguntas reais detectadas na base.
  3. Avalia as regras (quantas violações cada uma encontra) para o
     checkpoint de aprovação humana.
  4. Aplica a limpeza: apaga respostas em perguntas que deviam ter sido
     puladas, remove linhas sem nenhuma resposta e duplicados, e gera
     relatório detalhado do que foi alterado.

Não importa codificador.py (que exige OPENAI_API_KEY no import).
"""

from __future__ import annotations

import json
import os
import re
from io import BytesIO

import pandas as pd

MODELO_REGRAS = os.getenv("MODELO_REGRAS", "gpt-4o")

# Valores de célula considerados "vazios" na base
_VAZIOS = {"", "nan", "none", "-"}


# ─────────────────────────────────────────────────────────────────────────────
# 1. QUESTIONÁRIO (.docx) → TEXTO
# ─────────────────────────────────────────────────────────────────────────────

def extrair_texto_docx(data: bytes) -> str:
    """Extrai o texto do questionário (parágrafos e tabelas, na ordem)."""
    from docx import Document

    doc = Document(BytesIO(data))
    partes: list[str] = []
    for par in doc.paragraphs:
        txt = par.text.strip()
        if txt:
            partes.append(txt)
    for tabela in doc.tables:
        for linha in tabela.rows:
            celulas = [c.text.strip() for c in linha.cells if c.text.strip()]
            if celulas:
                partes.append(" | ".join(celulas))
    return "\n".join(partes)


# ─────────────────────────────────────────────────────────────────────────────
# 2. TEXTO + PERGUNTAS DA BASE → REGRAS ESTRUTURADAS (via IA)
# ─────────────────────────────────────────────────────────────────────────────

def _valores_observados(df: pd.DataFrame, pergunta: dict, max_vals: int = 12) -> list[str]:
    """Valores únicos observados nas colunas principais da pergunta."""
    vals: list[str] = []
    for col in pergunta.get("colunas", []):
        if col not in df.columns:
            continue
        serie = _serie(df, col).dropna().astype(str).str.strip()
        for v in serie.unique():
            if v.lower() not in _VAZIOS and v not in vals:
                vals.append(v)
        if len(vals) >= max_vals:
            break
    return vals[:max_vals]


def montar_resumo_perguntas(df: pd.DataFrame, perguntas: list[dict]) -> str:
    """Lista numerada das perguntas detectadas na base, com valores observados."""
    linhas: list[str] = []
    for p in perguntas:
        vals = _valores_observados(df, p)
        vals_txt = "; ".join(v[:60] for v in vals) if vals else "(sem valores)"
        linhas.append(
            f'{p["num"]} [{p["tipo"]}] "{p["pergunta"][:120]}"\n'
            f'    Valores observados: {vals_txt}'
        )
    return "\n".join(linhas)


_PROMPT_SISTEMA = """Você é um especialista em crítica e limpeza de bases de pesquisa quantitativa.
Sua tarefa é ler o texto de um questionário e extrair a LÓGICA DE PULOS, mapeando-a
para as perguntas reais da base de dados (P01, P02...).

Um "pulo" acontece quando a resposta a uma pergunta determina que o respondente NÃO
deve responder outra(s) pergunta(s). Exemplos no questionário:
  - "Se NÃO, pule para a P10"
  - "(SOMENTE PARA QUEM RESPONDEU SIM NA P3)"
  - "AGRADEÇA E ENCERRE A PESQUISA" (encerramento: todas as perguntas seguintes ficam vazias)

Responda SOMENTE com JSON válido, sem comentários, no formato:
{
  "regras": [
    {
      "tipo": "pulo",
      "descricao": "Quem respondeu 'Não' na P03 não responde P04 e P05",
      "se_pergunta": "P03",
      "se_valores": ["Não"],
      "perguntas_alvo": ["P04", "P05"]
    },
    {
      "tipo": "encerramento",
      "descricao": "Quem respondeu 'Não' na P01 encerra a pesquisa",
      "se_pergunta": "P01",
      "se_valores": ["Não (AGRADEÇA E ENCERRE A PESQUISA)"],
      "perguntas_alvo": []
    }
  ]
}

REGRAS OBRIGATÓRIAS:
1. Use SOMENTE os códigos P01, P02... da lista de perguntas da base fornecida.
2. Em "se_valores", use EXATAMENTE os textos dos "Valores observados" da base
   (não os textos do questionário, que podem ter grafia diferente).
3. Para "encerramento", deixe "perguntas_alvo" vazio — o sistema preenche com
   todas as perguntas seguintes.
4. Só crie regras que estejam explícitas ou claramente implícitas no questionário.
5. Se não houver nenhum pulo, retorne {"regras": []}.
"""


def _extrair_json(texto: str) -> dict:
    """Extrai o primeiro objeto JSON de uma resposta do modelo."""
    limpo = re.sub(r"```[a-z]*", "", texto).strip("`").strip()
    try:
        return json.loads(limpo)
    except json.JSONDecodeError:
        match = re.search(r"\{.*\}", limpo, re.DOTALL)
        if match:
            return json.loads(match.group())
        raise


def extrair_regras_questionario(
    texto_questionario: str,
    df: pd.DataFrame,
    perguntas: list[dict],
    modelo: str = None,
) -> list[dict]:
    """
    Chama a IA para extrair as regras de pulo do questionário.
    Retorna lista de regras validadas contra as perguntas da base.
    """
    from openai import OpenAI

    api_key = os.getenv("OPENAI_API_KEY", "")
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY não configurada.")

    client = OpenAI(api_key=api_key)
    resumo = montar_resumo_perguntas(df, perguntas)
    user_msg = (
        "PERGUNTAS DETECTADAS NA BASE:\n"
        f"{resumo}\n\n"
        "TEXTO DO QUESTIONÁRIO:\n"
        f"{texto_questionario[:30000]}"
    )

    resp = client.chat.completions.create(
        model=modelo or MODELO_REGRAS,
        messages=[
            {"role": "system", "content": _PROMPT_SISTEMA},
            {"role": "user", "content": user_msg},
        ],
        temperature=0,
    )
    dados = _extrair_json(resp.choices[0].message.content or "{}")
    return validar_regras(dados.get("regras", []), perguntas)


def validar_regras(regras: list[dict], perguntas: list[dict]) -> list[dict]:
    """
    Valida e normaliza as regras da IA:
      - descarta regras com pergunta condicional inexistente;
      - filtra alvos inexistentes;
      - expande encerramento para todas as perguntas seguintes;
      - adiciona id e flag 'ativa'.
    """
    nums = [p["num"] for p in perguntas]
    validas: list[dict] = []
    for i, r in enumerate(regras, start=1):
        se_p = str(r.get("se_pergunta", "")).strip()
        if se_p not in nums:
            continue
        valores = [str(v).strip() for v in r.get("se_valores", []) if str(v).strip()]
        if not valores:
            continue

        tipo = r.get("tipo", "pulo")
        alvo = [a for a in r.get("perguntas_alvo", []) if a in nums and a != se_p]
        if tipo == "encerramento":
            idx = nums.index(se_p)
            alvo = nums[idx + 1:]
        if not alvo:
            continue

        validas.append({
            "id": f"R{i}",
            "tipo": tipo,
            "descricao": str(r.get("descricao", "")).strip() or f"Regra {i}",
            "se_pergunta": se_p,
            "se_valores": valores,
            "perguntas_alvo": alvo,
            "ativa": True,
        })
    return validas


# ─────────────────────────────────────────────────────────────────────────────
# 3. AVALIAÇÃO E APLICAÇÃO DAS REGRAS
# ─────────────────────────────────────────────────────────────────────────────

def _norm(s: str) -> str:
    return re.sub(r"\s+", " ", str(s)).strip().casefold()


def _serie(df: pd.DataFrame, col: str) -> pd.Series:
    """Sempre retorna uma Series — primeira ocorrência se a coluna for duplicada."""
    s = df[col]
    if isinstance(s, pd.DataFrame):
        s = s.iloc[:, 0]
    return s


def _mascara_preenchida(df: pd.DataFrame, col: str) -> pd.Series:
    """Linhas em que a célula da coluna tem conteúdo de verdade."""
    s = _serie(df, col)
    return s.notna() & ~s.astype(str).str.strip().str.lower().isin(_VAZIOS)


def _pergunta_por_num(perguntas: list[dict], num: str) -> dict | None:
    for p in perguntas:
        if p["num"] == num:
            return p
    return None


def _cols_da_pergunta(pergunta: dict, df: pd.DataFrame) -> list[str]:
    """
    Todas as colunas da pergunta presentes na base (fechadas + 'Outro'),
    sem repetição — em perguntas 100% abertas, 'colunas' e 'cols_outros'
    são a mesma lista.
    """
    cols = dict.fromkeys(
        list(pergunta.get("colunas", [])) + list(pergunta.get("cols_outros", []))
    )
    return [c for c in cols if c in df.columns]


def _mascara_condicao(df: pd.DataFrame, pergunta: dict, valores: list[str]) -> pd.Series:
    """
    Linhas em que o respondente deu uma das respostas de `valores` na pergunta.
    Funciona para RU (valor na célula) e RM (célula da opção preenchida com o
    rótulo da opção). Match exato normalizado; se um valor não casar com nada,
    tenta prefixo (ex.: "Não" casa com "Não (AGRADEÇA E ENCERRE A PESQUISA)").
    """
    alvo_norm = [_norm(v) for v in valores]
    mascara = pd.Series(False, index=df.index)

    for col in pergunta.get("colunas", []):
        if col not in df.columns:
            continue
        serie = _serie(df, col).astype(str).map(_norm)
        casou = serie.isin(alvo_norm)
        if not casou.any():
            for v in alvo_norm:
                if len(v) >= 3:
                    casou = casou | serie.str.startswith(v)
        mascara = mascara | casou
    return mascara


def _mascara_respondida(df: pd.DataFrame, cols: list[str]) -> pd.Series:
    """Linhas com pelo menos uma resposta nas colunas dadas."""
    mascara = pd.Series(False, index=df.index)
    for col in dict.fromkeys(cols):
        mascara = mascara | _mascara_preenchida(df, col)
    return mascara


def avaliar_regras(df: pd.DataFrame, perguntas: list[dict],
                   regras: list[dict]) -> pd.DataFrame:
    """
    Para o checkpoint: quantas violações cada regra encontra na base.
    Violação = respondente que satisfaz a condição mas respondeu pergunta-alvo.
    """
    linhas = []
    for r in regras:
        p_cond = _pergunta_por_num(perguntas, r["se_pergunta"])
        if p_cond is None:
            continue
        cond = _mascara_condicao(df, p_cond, r["se_valores"])

        cols_alvo: list[str] = []
        for num in r["perguntas_alvo"]:
            p_alvo = _pergunta_por_num(perguntas, num)
            if p_alvo:
                cols_alvo += _cols_da_pergunta(p_alvo, df)

        violacao = cond & _mascara_respondida(df, cols_alvo)
        n_celulas = 0
        if violacao.any():
            for c in dict.fromkeys(cols_alvo):
                n_celulas += int((_mascara_preenchida(df, c) & violacao).sum())

        linhas.append({
            "id": r["id"],
            "Ativa": r.get("ativa", True),
            "Descrição": r["descricao"],
            "Condição": f'{r["se_pergunta"]} = {" | ".join(r["se_valores"])}',
            "Perguntas-alvo": ", ".join(r["perguntas_alvo"]),
            "Respondentes na condição": int(cond.sum()),
            "Violações (linhas)": int(violacao.sum()),
            "Células a limpar": n_celulas,
        })
    return pd.DataFrame(linhas)


def aplicar_limpeza(
    df: pd.DataFrame,
    perguntas: list[dict],
    regras: list[dict],
    remover_sem_resposta: bool = True,
    remover_duplicados: bool = True,
    col_id: str = "respondent_id",
) -> tuple[pd.DataFrame, pd.DataFrame, dict]:
    """
    Aplica as regras ativas e a limpeza genérica.

    Retorna (df_limpo, relatorio, resumo):
      relatorio — uma linha por ação: [respondente, acao, regra, detalhe]
      resumo    — contadores gerais para exibição.
    """
    df_limpo = df.copy()
    acoes: list[dict] = []

    ids = (_serie(df_limpo, col_id).astype(str) if col_id in df_limpo.columns
           else pd.Series(df_limpo.index.astype(str), index=df_limpo.index))

    # ── Regras de pulo: apaga respostas em perguntas que deviam ser puladas ──
    celulas_limpas = 0
    for r in regras:
        if not r.get("ativa", True):
            continue
        p_cond = _pergunta_por_num(perguntas, r["se_pergunta"])
        if p_cond is None:
            continue
        cond = _mascara_condicao(df_limpo, p_cond, r["se_valores"])

        cols_alvo: list[str] = []
        for num in r["perguntas_alvo"]:
            p_alvo = _pergunta_por_num(perguntas, num)
            if p_alvo:
                cols_alvo += _cols_da_pergunta(p_alvo, df_limpo)
        if not cols_alvo:
            continue

        respondida = _mascara_respondida(df_limpo, cols_alvo)
        violacao = cond & respondida
        cols_alvo = list(dict.fromkeys(cols_alvo))
        preench_por_col = {c: _mascara_preenchida(df_limpo, c) for c in cols_alvo}
        for idx in df_limpo.index[violacao]:
            preenchidas = [c for c in cols_alvo if preench_por_col[c].at[idx]]
            celulas_limpas += len(preenchidas)
            acoes.append({
                "respondente": ids.at[idx],
                "acao": "respostas removidas (pulo não respeitado)",
                "regra": f'{r["id"]} — {r["descricao"]}',
                "detalhe": f'{len(preenchidas)} célula(s): '
                           + "; ".join(c[:60] for c in preenchidas[:5])
                           + ("..." if len(preenchidas) > 5 else ""),
            })
            df_limpo.loc[idx, preenchidas] = pd.NA

    # ── Duplicados pelo id do respondente ────────────────────────────────────
    linhas_duplicadas = 0
    if remover_duplicados and col_id in df_limpo.columns:
        serie_id = _serie(df_limpo, col_id)
        dup = serie_id.duplicated(keep="first") & serie_id.notna()
        linhas_duplicadas = int(dup.sum())
        for idx in df_limpo.index[dup]:
            acoes.append({
                "respondente": ids.at[idx],
                "acao": "linha removida (respondente duplicado)",
                "regra": "limpeza genérica",
                "detalhe": f"{col_id} repetido",
            })
        df_limpo = df_limpo[~dup]

    # ── Linhas sem nenhuma resposta de pesquisa ──────────────────────────────
    linhas_vazias = 0
    if remover_sem_resposta:
        cols_perguntas: list[str] = []
        for p in perguntas:
            cols_perguntas += _cols_da_pergunta(p, df_limpo)
        if cols_perguntas:
            tem_resposta = _mascara_respondida(df_limpo, cols_perguntas)
            vazias = ~tem_resposta
            linhas_vazias = int(vazias.sum())
            for idx in df_limpo.index[vazias]:
                acoes.append({
                    "respondente": ids.at[idx],
                    "acao": "linha removida (nenhuma pergunta respondida)",
                    "regra": "limpeza genérica",
                    "detalhe": "",
                })
            df_limpo = df_limpo[tem_resposta]

    df_limpo = df_limpo.reset_index(drop=True)
    relatorio = pd.DataFrame(
        acoes, columns=["respondente", "acao", "regra", "detalhe"]
    )
    resumo = {
        "linhas_antes": len(df),
        "linhas_depois": len(df_limpo),
        "celulas_limpas": celulas_limpas,
        "linhas_vazias_removidas": linhas_vazias,
        "linhas_duplicadas_removidas": linhas_duplicadas,
    }
    return df_limpo, relatorio, resumo
